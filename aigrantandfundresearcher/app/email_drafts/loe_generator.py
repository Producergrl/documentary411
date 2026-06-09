"""
Letter of Enquiry (LOE) generator.
Produces a formatted .docx file per organization using Claude and python-docx.
Saved to ~/Documents/AIGrantAndFundResearcher/[FilmTitle]_Letters/
"""

import logging
import re
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

_MODEL = "claude-haiku-4-5-20251001"
_MAX_TOKENS = 1024

_SYSTEM_PROMPT = (
    "You write formal Letters of Enquiry for documentary filmmakers seeking grants "
    "and philanthropic funding. Your letters are professional, specific, compelling, "
    "and concise. Never use ellipses or hyphens as dashes. Use commas instead of "
    "hyphens where possible. Write in first person on behalf of the filmmaker. "
    "Do not use generic filler phrases. Every sentence must earn its place."
)


def generate_loe(
    cfg: dict,
    org_name: str,
    org_mission: str,
    contact_person: str,
    anthropic_api_key: str,
) -> str:
    """
    Generate the body text of an LOE via Claude.
    Returns the full letter body as a string.
    """
    full_name      = cfg.get("full_name", "")
    film_title     = cfg.get("film_title", "")
    logline        = cfg.get("logline", "")
    bio            = cfg.get("filmmaker_bio", "")
    total_budget   = cfg.get("total_budget", "")
    amount_seeking = cfg.get("amount_seeking", "")
    increments     = cfg.get("funding_increments", "")
    distribution   = cfg.get("distribution_plan", "")
    completion     = cfg.get("completion_date", "")

    salutation = (
        f"Dear {contact_person},"
        if contact_person and contact_person.strip()
        else f"Dear Grants Committee at {org_name},"
    )

    increment_line = (
        f" Contributions are also welcomed in increments of {increments}."
        if increments else ""
    )
    film_website  = cfg.get("film_website", "")
    website_line  = f" and website {film_website}" if film_website else ""

    prompt = f"""Write a formal Letter of Enquiry on behalf of {full_name} for the documentary film '{film_title}'.

Filmmaker bio: {bio}

Film logline: {logline}

About the organization being approached: {org_name}. Their mission: {org_mission}

Funding request: Total film budget is {total_budget}. Seeking {amount_seeking} to complete the film.{increment_line}

Distribution plan: {distribution}

Expected completion: {completion}

Write the letter in 5 paragraphs:
1. A strong opening that introduces the filmmaker and the film with authority, not humility.
2. A compelling description of the film, its investigation, and why it matters right now.
3. A specific paragraph connecting this film's mission directly to {org_name}'s stated mission — be precise, not generic.
4. The funding ask: what {amount_seeking} will accomplish, the distribution reach, and the community impact.
5. A confident closing inviting next steps, providing the filmmaker's contact information{website_line}.

Do not include the salutation or sign-off — those will be added separately.
Do not add headers or labels to each paragraph."""

    client = anthropic.Anthropic(api_key=anthropic_api_key.strip())
    message = client.messages.create(
        model=_MODEL,
        max_tokens=_MAX_TOKENS,
        system=_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": prompt}],
    )
    return salutation + "\n\n" + message.content[0].text.strip()


def save_loe_docx(
    cfg: dict,
    org_name: str,
    org_mission: str,
    contact_person: str,
    anthropic_api_key: str,
    output_folder: Path,
) -> Path:
    """
    Generate LOE text via Claude and save as a formatted .docx file.
    Returns the path to the saved file.
    """
    body_text = generate_loe(
        cfg=cfg,
        org_name=org_name,
        org_mission=org_mission,
        contact_person=contact_person,
        anthropic_api_key=anthropic_api_key,
    )

    doc = _build_docx(cfg, org_name, body_text)

    safe_org = re.sub(r"[^\w\s-]", "", org_name).strip().replace(" ", "_")[:50]
    filename = f"{safe_org}_LOE.docx"
    out_path = output_folder / filename
    doc.save(str(out_path))
    logger.info("Saved LOE: %s", out_path)
    return out_path


def _build_docx(cfg: dict, org_name: str, body_text: str) -> Document:
    """Build a cleanly formatted .docx document."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Header: Film title + filmmaker name
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(cfg.get("film_title", ""))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x18, 0x14)

    name_para = doc.add_paragraph()
    name_run = name_para.add_run(cfg.get("full_name", ""))
    name_run.font.size = Pt(10)
    name_run.font.color.rgb = RGBColor(0x6B, 0x63, 0x58)

    # Divider
    doc.add_paragraph("─" * 60)

    # Date + recipient
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph(org_name)
    doc.add_paragraph()

    # Body — split on double newlines into paragraphs
    paragraphs = [p.strip() for p in body_text.split("\n\n") if p.strip()]
    for i, para_text in enumerate(paragraphs):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(para_text)
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(10)

    # Sign-off
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig_run = sig.add_run(cfg.get("full_name", ""))
    sig_run.bold = True
    sig_run.font.size = Pt(11)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run("Writer, Producer, Director")
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = RGBColor(0x6B, 0x63, 0x58)

    contact_para = doc.add_paragraph()
    website      = cfg.get("film_website", "")
    contact_line = cfg.get("email", "")
    if website:
        contact_line = f"{contact_line}  |  {website}" if contact_line else website
    contact_run  = contact_para.add_run(contact_line)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(0x6B, 0x63, 0x58)

    return doc
